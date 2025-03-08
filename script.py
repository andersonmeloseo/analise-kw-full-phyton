import pandas as pd
import os
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.utils import get_column_letter
import time
from datetime import datetime
import difflib
import re
import numpy as np
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
import matplotlib.pyplot as plt
import xml.etree.ElementTree as ET
from xml.dom import minidom
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from openpyxl.chart import BarChart, PieChart, LineChart, Reference
from openpyxl.chart.label import DataLabelList

# =============================================================================
# Funções Auxiliares
# =============================================================================

def print_status(message):
    print(f"[STATUS] {message}")
    time.sleep(1)

def get_column_index(ws, column_name):
    for cell in ws[1]:
        if cell.value == column_name:
            return cell.column
    return None

def adjust_column_width(ws):
    for col in ws.columns:
        max_length = 0
        column = get_column_letter(col[0].column)
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = max_length + 2
        ws.column_dimensions[column].width = adjusted_width

def apply_header_style(ws):
    header_fill = PatternFill(start_color="000066", end_color="000066", fill_type="solid")
    header_font = Font(color="FFFFFF", size=12, bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

def apply_content_style(ws):
    content_font = Font(size=10)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = content_font

def apply_heatmap(ws, column_idx, values):
    values = pd.Series(values).dropna().sort_values()
    if values.empty:
        return

    total = len(values)
    red_threshold = int(total * 0.25)
    orange_threshold = int(total * 0.55)
    yellow_threshold = int(total * 0.85)

    red_max = values.iloc[red_threshold - 1] if red_threshold > 0 else values.min()
    orange_max = values.iloc[orange_threshold - 1] if orange_threshold > 0 else values.min()
    yellow_max = values.iloc[yellow_threshold - 1] if yellow_threshold > 0 else values.min()

    red_fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
    orange_fill = PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid")
    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    green_fill = PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")

    white_font = Font(color="FFFFFF", size=10)
    black_font = Font(color="000000", size=10)

    for row in ws.iter_rows(min_row=2, min_col=column_idx, max_col=column_idx):
        for cell in row:
            if cell.value is not None and not pd.isna(cell.value):
                val = float(cell.value)
                if val <= red_max:
                    cell.fill = red_fill
                    cell.font = white_font
                elif val <= orange_max:
                    cell.fill = orange_fill
                    cell.font = white_font
                elif val <= yellow_max:
                    cell.fill = yellow_fill
                    cell.font = black_font
                else:
                    cell.fill = green_fill
                    cell.font = black_font

def fuzzy_match(str1, str2, threshold=0.8):
    return difflib.SequenceMatcher(None, str1.lower(), str2.lower()).ratio() >= threshold

def carregar_cidades_brasil():
    try:
        cidades_df = pd.read_excel('cidades_brasil.xlsx')
        return cidades_df['CIDADE'].str.lower().str.strip().tolist()
    except FileNotFoundError:
        print_status("Erro: Arquivo 'cidades_brasil.xlsx' não encontrado na pasta do script!")
        raise
    except KeyError:
        print_status("Erro: Coluna 'CIDADE' não encontrada na planilha 'cidades_brasil.xlsx'!")
        raise
    except Exception as e:
        print_status(f"Erro ao carregar cidades: {str(e)}")
        raise

def mapear_objetivo(intent):
    intent = str(intent).lower()
    if "informational" in intent:
        return "Mais Acessos"
    elif "transactional" in intent or "transacional" in intent:
        return "Vendas no E-commerce"
    elif "commercial" in intent:
        return "Captura de Leads"
    elif "navegacional" in intent:
        return "Branding/Autoridade"
    else:
        return "Outro"

# Funções para o Relatório
def add_title(doc, text):
    title = doc.add_heading(text, level=1)
    title.style.font.name = 'Arial'
    title.style.font.size = Pt(16)
    title.style.font.color.rgb = RGBColor(0, 0, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_subtitle(doc, text):
    subtitle = doc.add_heading(text, level=2)
    subtitle.style.font.name = 'Arial'
    subtitle.style.font.size = Pt(14)
    subtitle.style.font.color.rgb = RGBColor(0, 102, 204)

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_image(doc, image_path, width=Inches(5)):
    doc.add_picture(image_path, width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# =============================================================================
# Mapeamento da Jornada e Tipologia
# =============================================================================

def get_etapa_da_jornada(intent):
    intent = str(intent).strip().lower()
    if intent == "informational":
        return "Conscientização"
    elif intent in ["transactional", "transacional"]:
        return "Decisão"
    elif intent == "commercial":
        return "Consideração"
    elif intent == "navegacional":
        return "Fidelização"
    else:
        return "Sem Jornada Definida"

def get_tipologia_sugerida(row):
    intent = str(row.get('Intent', '')).strip().lower()
    serp = str(row.get('SERP Features', '')).strip().lower()

    if intent == "informational":
        if "featured snippets" in serp:
            return "Artigo de Blog"
        elif "instant answer" in serp:
            return "Artigo de Blog (respostas rápidas)"
        elif any(term in serp for term in ["video", "featured video", "video carousel"]):
            return "Guia"
        elif any(term in serp for term in ["image", "image pack"]):
            return "Infográfico"
        elif "people also ask" in serp:
            return "FAQs"
        elif "knowledge panel" in serp:
            return "Artigo de Blog (definições amplas)"
        elif any(term in serp for term in ["news", "top stories"]):
            return "Notícias do Setor"
        else:
            return "Análise Manual"
    elif intent in ["transactional", "transacional"]:
        if any(term in serp for term in ["shopping ads", "ads top", "ads bottom", "ads middle"]):
            return "Página de Produto/Serviço (otimizadas para conversão)"
        elif any(term in serp for term in ["hotel pack", "flights", "recipes", "jobs"]):
            return "Página de Produto/Serviço"
        elif "buying guide" in serp:
            return "Página de Produto/Serviço"
        elif any(term in serp for term in ["popular products", "related products", "organic carousel"]):
            return "Página de Produto/Serviço"
        elif any(term in serp for term in ["address pack", "twitter carousel"]):
            return "Página de Produto/Serviço"
        else:
            return "Análise Manual"
    elif intent == "commercial":
        if any(term in serp for term in ["featured reviews", "video carousel"]):
            return "Comparativo"
        elif "buying guide" in serp:
            return "Comparativo"
        elif "discussions and forums" in serp:
            return "Comparativo"
        elif any(term in serp for term in ["brands", "explore", "related searches", "related products"]):
            return "Comparativo"
        elif "questions and answers" in serp:
            return "FAQs"
        else:
            return "Análise Manual"
    elif intent == "navegacional":
        if "sitelinks" in serp:
            return "Documentação"
        elif "knowledge panel" in serp:
            return "Blog de Atualizações"
        elif any(term in serp for term in ["twitter", "twitter carousel"]):
            return "Blog de Atualizações"
        elif any(term in serp for term in ["find results on", "address pack"]):
            return "Página de Suporte"
        else:
            return "Análise Manual"
    else:
        return "Análise Manual"

# =============================================================================
# Função para Estratégia com Palavras-Chave
# =============================================================================

def criar_planilha_palavras_por_estrategia(folder_name, objective, combined_df):
    print_status("Criando a planilha 'Palavras por Estratégia.xlsx'...")

    dados_estrategia = [
        ["Captura de Leads", "Local SEO", "Melhor [serviço] em [cidade]", "Local Pack, Snippets", "Landing Page, Blog Local", "SEO Local + Conversão"],
        ["Vendas no E-commerce", "Transactional", "Comprar [produto] com desconto", "Shopping Ads, Reviews", "Páginas de Produto, Comparativos", "CRO + SEO para Produtos"],
        ["Mais Acessos", "Informational", "Como funciona [assunto]", "Featured Snippets, PAA", "Blog, Guia Completo", "SEO para Topo de Funil"],
        ["Monetização com Adsense", "High CPC", "Melhor seguro de saúde nos EUA", "Featured Snippets, PAA", "Lista Comparativa, Blog", "SEO para Alto CPC"],
        ["Branding/Autoridade", "Institucional", "[Empresa] é confiável?", "Knowledge Panel, Twitter Carousel", "Página Institucional, Blog", "SEO para Reputação"],
        ["Outro", "Custom", "Personalizado conforme análise", "Variável", "Variável", "Definido pelo usuário"]
    ]

    estrategia_base_df = pd.DataFrame(dados_estrategia, columns=["Objetivo", "Keyword Type", "Exemplo de Palavra-chave", "SERP Features", "Tipologia de Conteúdo", "Estratégia"])

    objetivo_map = {
        "1": "Captura de Leads",
        "2": "Vendas no E-commerce",
        "3": "Mais Acessos",
        "4": "Monetização com Adsense",
        "5": "Branding/Autoridade",
        "6": "Outro"
    }
    objetivo_selecionado = objetivo_map.get(objective, "Outro") if objective in "123456" else objective.capitalize()

    estrategia_df = combined_df.copy()
    estrategia_df["Objetivo"] = estrategia_df["Intent"].apply(mapear_objetivo)
    estrategia_df = estrategia_df.merge(estrategia_base_df.drop(columns=["Exemplo de Palavra-chave"]),
                                        on="Objetivo", how="left", suffixes=('', '_base'))

    estrategia_df = estrategia_df.rename(columns={"Keyword": "Palavra-chave"})
    colunas_finais = ["Objetivo", "Palavra-chave", "Volume", "Intent", "SERP Features",
                      "Keyword Type", "Tipologia de Conteúdo", "Estratégia"]
    estrategia_df = estrategia_df[colunas_finais]

    if objetivo_selecionado != "Outro":
        estrategia_df = estrategia_df[estrategia_df["Objetivo"] == objetivo_selecionado]

    wb = Workbook()
    ws = wb.active
    ws.title = "Palavras por Estratégia"
    for r in dataframe_to_rows(estrategia_df, index=False, header=True):
        ws.append(r)

    volume_col_index = get_column_index(ws, "Volume")
    if volume_col_index and not estrategia_df['Volume'].dropna().empty:
        apply_heatmap(ws, volume_col_index, estrategia_df['Volume'])

    apply_header_style(ws)
    apply_content_style(ws)
    adjust_column_width(ws)
    wb.save(os.path.join(folder_name, "Palavras por Estratégia.xlsx"))
    print_status("Planilha 'Palavras por Estratégia.xlsx' criada com sucesso!")
    return estrategia_df, objetivo_selecionado

# =============================================================================
# Função para Planejamento de Crescimento
# =============================================================================

def criar_planilha_planejamento_crescimento(folder_name, combined_df, volume_atual, crescimento_mensal, meses_planejamento, palavras_por_mes, objective, cidades_brasil):
    print_status("Criando a planilha 'Planejamento de Crescimento.xlsx'...")
    output_path = os.path.join(folder_name, "Planejamento de Crescimento.xlsx")

    ctr_rates = {
        1: (0.25, 0.35), 2: (0.15, 0.20), 3: (0.10, 0.15), 4: (0.07, 0.10), 5: (0.05, 0.07),
        6: (0.04, 0.06), 7: (0.03, 0.05), 8: (0.02, 0.04), 9: (0.02, 0.03), 10: (0.01, 0.02)
    }

    crescimento_absoluto = volume_atual * (crescimento_mensal / 100)
    acessos_alvo = volume_atual + crescimento_absoluto
    volume_min_por_palavra = crescimento_absoluto / (palavras_por_mes * ctr_rates[10][0])
    volume_max_por_palavra = crescimento_absoluto / (palavras_por_mes * ctr_rates[1][1])

    dados_calculo = [
        ["Volume Atual (mensal)", volume_atual],
        ["Crescimento Desejado (%)", crescimento_mensal],
        ["Crescimento Absoluto (acessos)", crescimento_absoluto],
        ["Acessos Alvo (mensal)", acessos_alvo],
        ["Meses de Planejamento", meses_planejamento],
        ["Palavras por Mês", palavras_por_mes],
        ["Volume Mínimo por Palavra (pior cenário - Posição 10)", round(volume_min_por_palavra)],
        ["Volume Máximo por Palavra (melhor cenário - Posição 1)", round(volume_max_por_palavra)]
    ]
    calculo_df = pd.DataFrame(dados_calculo, columns=["Métrica", "Valor"])

    objetivo_map = {
        "1": "Captura de Leads",
        "2": "Vendas no E-commerce",
        "3": "Mais Acessos",
        "4": "Monetização com Adsense",
        "5": "Branding/Autoridade",
        "6": "Outro"
    }
    objetivo_selecionado = objetivo_map.get(objective, "Outro") if objective in "123456" else objective.capitalize()

    palavras_df = combined_df.copy()
    palavras_df["Objetivo"] = palavras_df["Intent"].apply(mapear_objetivo)
    if objetivo_selecionado != "Outro":
        palavras_df = palavras_df[palavras_df["Objetivo"] == objetivo_selecionado]

    palavras_df["Keyword"] = palavras_df["Keyword"].fillna("").astype(str)

    # Filtro de cidades menos restritivo (apenas palavras exatas de cidades)
    palavras_df = palavras_df[~palavras_df["Keyword"].str.lower().apply(
        lambda x: any(cidade in x.split() for cidade in cidades_brasil)
    )]

    print_status(f"Linhas após filtro de cidades: {len(palavras_df)}")
    if palavras_df.empty:
        print_status("Aviso: O DataFrame está vazio após o filtro de cidades!")
        return None, None, None, None, None, None, None, None

    volume_col = None
    for col in palavras_df.columns:
        if "volume" in col.lower():
            volume_col = col
            break
    if not volume_col:
        print_status("Erro: Nenhuma coluna de volume encontrada no DataFrame!")
        return None, None, None, None, None, None, None, None

    sort_columns = [volume_col]
    if "Competitive Density" in palavras_df.columns:
        sort_columns.append("Competitive Density")
        palavras_df = palavras_df.sort_values(by=sort_columns, ascending=[False, True])
    else:
        palavras_df = palavras_df.sort_values(by=volume_col, ascending=False)

    total_palavras = min(meses_planejamento * palavras_por_mes, len(palavras_df))
    palavras_selecionadas = palavras_df.head(total_palavras)
    colunas_selecao = ["Keyword", volume_col, "Intent", "SERP Features"] + \
                     (["Competitive Density"] if "Competitive Density" in palavras_df.columns else [])

    palavras_por_mes_ajustado = max(1, total_palavras // meses_planejamento)
    meses = []
    for i in range(meses_planejamento):
        inicio = i * palavras_por_mes_ajustado
        fim = min((i + 1) * palavras_por_mes_ajustado, total_palavras)
        mes_df = palavras_selecionadas.iloc[inicio:fim][colunas_selecao]
        meses.append(mes_df)

    # Segmentação por comprimento
    palavras_df["Comprimento"] = palavras_df["Keyword"].apply(lambda x: len(str(x).strip().split()))
    cauda_curta = palavras_df[palavras_df["Comprimento"] <= 2][colunas_selecao]
    cauda_media = palavras_df[palavras_df["Comprimento"] == 3][colunas_selecao]
    cauda_longa = palavras_df[palavras_df["Comprimento"] >= 4][colunas_selecao]

    # Grupos Semânticos
    total_grupos = max(2, palavras_por_mes // 2)  # Garantir pelo menos 2 grupos
    total_palavras_semantico = min(len(palavras_df), total_grupos * 20)  # Aumentar para 20 por grupo
    palavras_semantico = palavras_df.head(total_palavras_semantico).copy()  # Criar uma cópia explícita
    if len(palavras_semantico) >= 10:
        vectorizer = TfidfVectorizer(max_features=5000, stop_words=None)
        X = vectorizer.fit_transform(palavras_semantico["Keyword"])
        kmeans = KMeans(n_clusters=total_grupos, n_init=10, random_state=42)  # Adicionar n_init=10
        palavras_semantico["Grupo Semântico"] = kmeans.fit_predict(X).astype(str)  # Usar atribuição direta
        palavras_semantico["Grupo Semântico"] = "Grupo_" + palavras_semantico["Grupo Semântico"]
    else:
        palavras_semantico["Grupo Semântico"] = "Sem Grupo (poucas palavras)"
    colunas_semantico = ["Grupo Semântico"] + colunas_selecao

    # Palavras para Blog (informacionais)
    palavras_blog = palavras_df[palavras_df["Intent"].str.lower().str.contains("informational", na=False)][colunas_selecao]
    if palavras_blog.empty:
        print_status("Aviso: Nenhuma palavra informacional encontrada para blog!")

    # Criação da planilha
    wb = Workbook()

    ws_calculo = wb.active
    ws_calculo.title = "Calculo de Crescimento"
    for r in dataframe_to_rows(calculo_df, index=False, header=True):
        ws_calculo.append(r)
    apply_header_style(ws_calculo)
    apply_content_style(ws_calculo)
    adjust_column_width(ws_calculo)

    ws_selecao = wb.create_sheet("Selecao de Palavras")
    headers = ["Mês"] + colunas_selecao
    ws_selecao.append(headers)
    for i, mes_df in enumerate(meses, 1):
        for idx, row in mes_df.iterrows():
            ws_selecao.append([f"Mês {i}"] + row.tolist())
    volume_col_index = get_column_index(ws_selecao, volume_col)
    if volume_col_index and not palavras_selecionadas[volume_col].dropna().empty:
        apply_heatmap(ws_selecao, volume_col_index, palavras_selecionadas[volume_col])
    apply_header_style(ws_selecao)
    apply_content_style(ws_selecao)
    adjust_column_width(ws_selecao)

    ws_curta = wb.create_sheet("Cauda Curta")
    for r in dataframe_to_rows(cauda_curta, index=False, header=True):
        ws_curta.append(r)
    volume_col_index = get_column_index(ws_curta, volume_col)
    if volume_col_index and not cauda_curta[volume_col].dropna().empty:
        apply_heatmap(ws_curta, volume_col_index, cauda_curta[volume_col])
    apply_header_style(ws_curta)
    apply_content_style(ws_curta)
    adjust_column_width(ws_curta)

    ws_media = wb.create_sheet("Cauda Media")
    for r in dataframe_to_rows(cauda_media, index=False, header=True):
        ws_media.append(r)
    volume_col_index = get_column_index(ws_media, volume_col)
    if volume_col_index and not cauda_media[volume_col].dropna().empty:
        apply_heatmap(ws_media, volume_col_index, cauda_media[volume_col])
    apply_header_style(ws_media)
    apply_content_style(ws_media)
    adjust_column_width(ws_media)

    ws_longa = wb.create_sheet("Cauda Longa")
    for r in dataframe_to_rows(cauda_longa, index=False, header=True):
        ws_longa.append(r)
    volume_col_index = get_column_index(ws_longa, volume_col)
    if volume_col_index and not cauda_longa[volume_col].dropna().empty:
        apply_heatmap(ws_longa, volume_col_index, cauda_longa[volume_col])
    apply_header_style(ws_longa)
    apply_content_style(ws_longa)
    adjust_column_width(ws_longa)

    ws_semantico = wb.create_sheet("Grupos Semanticos")
    for r in dataframe_to_rows(palavras_semantico[colunas_semantico], index=False, header=True):
        ws_semantico.append(r)
    volume_col_index = get_column_index(ws_semantico, volume_col)
    if volume_col_index and not palavras_semantico[volume_col].dropna().empty:
        apply_heatmap(ws_semantico, volume_col_index, palavras_semantico[volume_col])
    apply_header_style(ws_semantico)
    apply_content_style(ws_semantico)
    adjust_column_width(ws_semantico)

    ws_blog = wb.create_sheet("Palavras para Blog")
    for r in dataframe_to_rows(palavras_blog, index=False, header=True):
        ws_blog.append(r)
    volume_col_index = get_column_index(ws_blog, volume_col)
    if volume_col_index and not palavras_blog[volume_col].dropna().empty:
        apply_heatmap(ws_blog, volume_col_index, palavras_blog[volume_col])
    apply_header_style(ws_blog)
    apply_content_style(ws_blog)
    adjust_column_width(ws_blog)

    wb.save(output_path)
    print_status("Planilha 'Planejamento de Crescimento.xlsx' criada com sucesso!")
    return calculo_df, palavras_selecionadas, cauda_curta, cauda_media, cauda_longa, palavras_semantico, palavras_blog, meses, colunas_selecao

# =============================================================================
# Função para Top 100 Palavras-Chave por Tipo
# =============================================================================

def criar_planilha_top_palavras_por_tipo(folder_name, combined_df):
    print_status("Criando a planilha 'Top 100 Palavras por Tipo.xlsx'...")

    # Identificar a coluna de volume dinamicamente
    volume_col = None
    for col in combined_df.columns:
        if "volume" in col.lower():
            volume_col = col
            break
    if not volume_col:
        print_status("Erro: Nenhuma coluna de volume encontrada no DataFrame!")
        return None

    # Garantir que "Keyword" seja string e tratar valores nulos
    combined_df["Keyword"] = combined_df["Keyword"].fillna("").astype(str)

    # Filtrar palavras com volume válido
    df_valid = combined_df[combined_df[volume_col].notna() & (combined_df[volume_col] > 0)].copy().reset_index(drop=True)
    if len(df_valid) < 1:
        print_status("Erro: Nenhuma palavra com volume válido encontrada!")
        return None

    # Verificar se existe uma coluna para agrupar tipos (ex.: "Keyword Type", "Categoria", etc.)
    type_col = None
    for col in df_valid.columns:
        if "type" in col.lower() or "categoria" in col.lower() or "grupo" in col.lower():
            type_col = col
            break

    top_palavras = {}
    if type_col:
        # Agrupar por tipo fornecido na coluna identificada
        print_status(f"Agrupando palavras por '{type_col}' fornecido na planilha...")
        for tipo in df_valid[type_col].dropna().unique():
            tipo_df = df_valid[df_valid[type_col] == tipo].sort_values(by=volume_col, ascending=False)
            if not tipo_df.empty:
                # Garantir que o nome do tipo tenha menos de 31 caracteres (limite do Excel)
                tipo_name = str(tipo)[:31]
                top_palavras[tipo_name] = tipo_df.head(100)[["Keyword", volume_col, "Intent"]]
    else:
        # Fallback: usar clustering TF-IDF se não houver coluna de tipo
        print_status("Nenhuma coluna de tipo encontrada. Usando clustering automático como fallback...")
        if len(df_valid) >= 10:
            vectorizer = TfidfVectorizer(max_features=5000, stop_words=None)
            X = vectorizer.fit_transform(df_valid["Keyword"])
            feature_names = vectorizer.get_feature_names_out()
            n_clusters = min(10, max(2, len(df_valid) // 100))
            kmeans = KMeans(n_clusters=n_clusters, n_init=10, random_state=42)
            df_valid["Cluster"] = kmeans.fit_predict(X)

            for cluster_id in range(n_clusters):
                cluster_df = df_valid[df_valid["Cluster"] == cluster_id].sort_values(by=volume_col, ascending=False)
                if not cluster_df.empty:
                    cluster_indices = df_valid.index[df_valid["Cluster"] == cluster_id].tolist()
                    cluster_tfidf = X[cluster_indices].mean(axis=0).A1
                    top_indices = cluster_tfidf.argsort()[-2:][::-1]
                    top_terms = [feature_names[idx] for idx in top_indices]
                    cluster_name = " ".join(top_terms).capitalize()[:31]
                    if cluster_name in top_palavras:
                        cluster_name = f"{cluster_name} {cluster_id}"[:31]
                    top_palavras[cluster_name] = cluster_df.head(100)[["Keyword", volume_col, "Intent"]]
        else:
            print_status("Menos de 10 palavras válidas. Agrupando todas em 'Geral'...")
            top_palavras["Geral"] = df_valid.sort_values(by=volume_col, ascending=False).head(100)[["Keyword", volume_col, "Intent"]]

    # Criar a planilha com os grupos
    wb = Workbook()
    for tipo, df in top_palavras.items():
        ws = wb.create_sheet(tipo)  # Nome da aba será o tipo identificado
        for r in dataframe_to_rows(df, index=False, header=True):
            ws.append(r)
        volume_col_index = get_column_index(ws, volume_col)
        if volume_col_index and not df[volume_col].dropna().empty:
            apply_heatmap(ws, volume_col_index, df[volume_col])
        apply_header_style(ws)
        apply_content_style(ws)
        adjust_column_width(ws)

    if "Sheet" in wb.sheetnames:
        wb.remove(wb["Sheet"])

    wb.save(os.path.join(folder_name, "Top 100 Palavras por Tipo.xlsx"))
    print_status("Planilha 'Top 100 Palavras por Tipo.xlsx' criada com sucesso!")
    return top_palavras
# =============================================================================
# Função para Palavras para Ads Filtradas (com KW Negativas Excluídas)
# =============================================================================

def criar_planilha_palavras_para_ads_filtradas(folder_name, combined_df):
    print_status("Criando a planilha 'Palavras para Ads Filtradas.xlsx'...")

    # Identificar a coluna de volume
    volume_col = None
    for col in combined_df.columns:
        if "volume" in col.lower():
            volume_col = col
            break
    if not volume_col:
        print_status("Erro: Nenhuma coluna de volume encontrada no DataFrame!")
        return None, None

    # Lista de palavras negativas fornecida
    kw_negativas = [
        "joguinho", "jogar", "reclamação", "reclamacoes", "reclamaçao", "reclamaçoes", "reclamação", "reclamações",
        "problema", "problemas", "tragedia", "olx", "tragedias", "tragédia", "reclameaqui", "acidente", "acidentes",
        "devolucao", "devolucoes", "devoluçao", "devoluçoes", "devolução", "devoluções", "falso", "falsos",
        "falsificado", "falsificados", "gratis", "grátis", "gratuito", "gratuitos", "de graça", "proibido",
        "proibidos", "paraguay", "paraguai", "download", "baixar", "receita", "receitas", "ilegal", "quero ver",
        "o que é", "destilaria", "testosterona", "mal", "mau", "vaga", "vagas", "emprego", "empregos",
        "agência de empregos", "agência de emprego", "agencia de emprego", "agencie de empregos", "quanto ganha",
        "curriculo", "currículo", "agência de empregos", "oferta de emprego", "mercadolivre", "mercado livre",
        "acompanhantes", "afundar", "afundou", "aids", "boate", "crime", "drogas", "escravo", "hospital",
        "incêndio", "incendio", "miséria", "naufragio", "naufrágio", "naufrágios", "pobreza", "pornô", "hentai",
        "pdf", "pornozão", "pornografia", "prostituição", "sexo", "sexual", "sexy", "torrent", "violência",
        "xvideos", "x videos", "airbnb", "tragédias", "jogo", "jogos", "defeito", "defeitos", "free", "game",
        "games", "joguinhos", "VX case", "lente de contato", "wallpaper", "papel de parede", "Sailor", "Ramen",
        "pior", "queixa", "fraude", "escândalo", "negativo", "lei", "legal", "legislação", "regra", "regulação",
        "regulamento", "faq", "termos", "condições", "perguntas frequentes", "politica de privacidade", "ajuda",
        "condições de venda", "guia de tamanhos", "estorno", "rejeição", "email", "garantia", "manual", "recolha",
        "reembolso", "substituir", "substituição", "rma", "apoio", "amostras", "leilão", "passatempo", "revenda",
        "segunda mão", "usado", "pode", "como", "o que", "quando", "onde", "quem", "porque", "carreira", "estágio",
        "trabalho", "trabalhar", "recursos humanos", "gestor rh", "recrutamento", "cv", "freelancer", "aulas",
        "cursos", "formação", "escola", "treino", "universidade", "faculdade", "especialização", "graduação",
        "mestrado", "doutorado", "associação", "jornal", "revista", "métricas", "notícias", "investigação",
        "review", "opinião", "estatísticas", "histórias", "tutorial", "definição", "significado", "sobre",
        "relatorios", "especificações", "behance", "facebook", "flickr", "instagram", "linkedin", "meetup",
        "messenger", "pinterest", "reddit", "snapchat", "soundcloud", "telegram", "tiktok", "tumblr", "twitter",
        "vimeo", "valor", "salário", "média salarial", "o q é", "dói", "doi", "dor", "doer", "insuportável",
        "quanto", "o que faz", "como faz", "como é", "apresentação", "presencial", "nome para", "quanto fatura",
        "dá dinheiro", "oq faz", "0800", "de graca", "graça", "gratiz", "grátiz", "gratuita", "sem", "sem custo",
        "sem pagar", "gratúitos", "gratúitas", "gratuítos", "gratuítas", "o que e", "oq e", "o q e", "pdf.",
        "proposta comercial", "modelo", "frase", "frases", "foto", "fotos", "imagem", "imagens", "fotografia",
        "fotografias", "vídeo", "video", "vídeos", "videos", "dica", "dicas", "antes", "depois", "antes e depois",
        "www", "digulgação", "divulgacao", "divulgaçao", "divulgacão", "Custo zero", "Download gratuito",
        "Versão gratuita", "Trial gratuito", "Demonstração gratuita", "Teste grátis", "Experimente grátis",
        "Experimentar grátis", "catho", "cathu", "cato", "catu", "Consultora comercial", "Consultora de vendas",
        "Consultores", "contratação", "curriculu", "curriculum", "curriculun", "entrevista", "estagios",
        "estágios", "infojob", "infojobs", "jovem aprendis", "jovem aprendiz", "labuta", "Lista", "manager",
        "ocupação", "oportunidade", "rio vagas", "riovagas", "Salario", "Salário", "Sandra mara", "sandramara",
        "serviço", "servisso", "Telemarketing", "vitae", "Contrata", "Trabalhe", "Trabalhe conosco", "servico",
        "Ganha", "Conosco", "Contratar", "auxiliar", "diretor", "supervisor", "gerente", "função", "funções",
        "cargo", "cargos"
    ]
    # Remover duplicatas da lista
    kw_negativas = list(set(kw_negativas))
    print_status(f"{len(kw_negativas)} palavras negativas carregadas da lista interna")

    # Garantir que "Keyword" seja string e tratar valores nulos
    combined_df["Keyword"] = combined_df["Keyword"].fillna("").astype(str)

    # Criar uma cópia do DataFrame original para trabalhar
    df_inicial = combined_df.copy()

    # Filtrar palavras que NÃO contenham termos negativos (Palavras Filtradas)
    palavras_ads_filtradas = df_inicial[~df_inicial["Keyword"].str.lower().apply(
        lambda x: any(negativa in x.lower() for negativa in kw_negativas)
    )]

    # Identificar palavras excluídas da lista inicial que CONTENHAM termos negativos
    palavras_excluidas = df_inicial[df_inicial["Keyword"].str.lower().apply(
        lambda x: any(negativa in x.lower() for negativa in kw_negativas)
    )]

    print_status(f"Palavras filtradas (não contêm negativas): {len(palavras_ads_filtradas)}")
    print_status(f"Palavras excluídas negativadas (excluídas da lista inicial): {len(palavras_excluidas)}")

    # Se não houver palavras filtradas, criar DataFrame vazio
    if palavras_ads_filtradas.empty:
        print_status("Aviso: Nenhuma palavra restante após exclusão de negativas!")
        palavras_ads_filtradas = pd.DataFrame(columns=["Keyword", volume_col, "Intent", "SERP Features"])

    # Se não houver palavras excluídas, criar DataFrame vazio
    if palavras_excluidas.empty:
        print_status("Aviso: Nenhuma palavra excluída encontrada na lista inicial!")
        palavras_excluidas = pd.DataFrame(columns=["Keyword", volume_col, "Intent", "SERP Features"])

    # Ordenar por volume (se disponível)
    if volume_col in palavras_ads_filtradas.columns and not palavras_ads_filtradas[volume_col].dropna().empty:
        palavras_ads_filtradas = palavras_ads_filtradas.sort_values(by=volume_col, ascending=False)
    else:
        palavras_ads_filtradas = palavras_ads_filtradas.sort_values(by="Keyword")

    if volume_col in palavras_excluidas.columns and not palavras_excluidas[volume_col].dropna().empty:
        palavras_excluidas = palavras_excluidas.sort_values(by=volume_col, ascending=False)
    else:
        palavras_excluidas = palavras_excluidas.sort_values(by="Keyword")

    # Selecionar colunas relevantes
    colunas_selecao = ["Keyword", volume_col, "Intent", "SERP Features"] if volume_col else ["Keyword", "Intent", "SERP Features"]

    # Criar a planilha com duas abas
    wb = Workbook()

    # Aba 1: Palavras Filtradas
    ws_filtradas = wb.active
    ws_filtradas.title = "Palavras Filtradas"
    for r in dataframe_to_rows(palavras_ads_filtradas[colunas_selecao], index=False, header=True):
        ws_filtradas.append(r)
    volume_col_index = get_column_index(ws_filtradas, volume_col) if volume_col else None
    if volume_col_index and not palavras_ads_filtradas[volume_col].dropna().empty:
        apply_heatmap(ws_filtradas, volume_col_index, palavras_ads_filtradas[volume_col])
    apply_header_style(ws_filtradas)
    apply_content_style(ws_filtradas)
    adjust_column_width(ws_filtradas)

    # Aba 2: Palavras Excluídas Negativadas
    ws_excluidas = wb.create_sheet("Palavras Excluidas Negativadas")
    for r in dataframe_to_rows(palavras_excluidas[colunas_selecao], index=False, header=True):
        ws_excluidas.append(r)
    volume_col_index = get_column_index(ws_excluidas, volume_col) if volume_col else None
    if volume_col_index and not palavras_excluidas[volume_col].dropna().empty:
        apply_heatmap(ws_excluidas, volume_col_index, palavras_excluidas[volume_col])
    apply_header_style(ws_excluidas)
    apply_content_style(ws_excluidas)
    adjust_column_width(ws_excluidas)

    # Salvar a planilha
    output_path = os.path.join(folder_name, "Palavras para Ads Filtradas.xlsx")
    wb.save(output_path)
    print_status("Planilha 'Palavras para Ads Filtradas.xlsx' criada com sucesso!")
    return palavras_ads_filtradas, palavras_excluidas

# =============================================================================
# Função para Criar Dashboard Profissional
# =============================================================================

def criar_dashboard_profissional(folder_name, combined_df, intent_counts, serp_counts, jornada_counts, ctr_export_df, estrategia_df, calculo_df, meses, volume_col, objective, now, ctr_rates):
    print_status("Criando Dashboard Profissional no Excel...")

    wb = Workbook()
    ws = wb.active
    ws.title = "Dashboard"

    # Estilo para bordas e preenchimento
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    header_fill = PatternFill(start_color="000066", end_color="000066", fill_type="solid")
    section_fill = PatternFill(start_color="E6F0FA", end_color="E6F0FA", fill_type="solid")

    # --- Cabeçalho ---
    ws['A1'] = "Dashboard de Análise de Palavras-Chave para SEO"
    ws['A1'].font = Font(size=16, bold=True, color="FFFFFF")
    ws['A1'].fill = header_fill
    ws.merge_cells('A1:H1')
    ws['A1'].alignment = Alignment(horizontal="center", vertical="center")

    ws['A2'] = f"Data: {now.strftime('%d-%m-%Y %H:%M:%S')} | Objetivo: {objective}"
    ws['A2'].font = Font(size=12, italic=True)
    ws.merge_cells('A2:H2')

    # --- Seção 1: Resumo Geral ---
    ws['A4'] = "Resumo Geral"
    ws['A4'].font = Font(size=14, bold=True)
    ws['A4'].fill = section_fill
    resumo_data = [
        ["Total de Palavras", len(combined_df)],
        ["Volume Médio", round(combined_df[volume_col].mean(), 2) if volume_col and not combined_df[volume_col].dropna().empty else "N/A"],
        ["Palavras com Intent", len(combined_df[combined_df['Intent'].notna()])]
    ]
    for i, (label, value) in enumerate(resumo_data, start=5):
        ws[f'A{i}'] = label
        ws[f'B{i}'] = value
        ws[f'A{i}'].font = Font(bold=True)
        ws[f'B{i}'].alignment = Alignment(horizontal="center")
        for col in ['A', 'B']:
            ws[f'{col}{i}'].border = thin_border

    # --- Seção 2: Distribuição por Intenção ---
    ws['A10'] = "Distribuição por Intenção de Busca"
    ws['A10'].font = Font(size=14, bold=True)
    ws['A10'].fill = section_fill
    intents = ['Informational', 'Transactional', 'Commercial', 'Navigational']
    ws.append(["Intenção", "Quantidade"])
    for intent in intents:
        ws.append([intent, intent_counts.get(intent, 0)])
    for row in range(11, 16):
        for col in ['A', 'B']:
            ws[f'{col}{row}'].border = thin_border
    ws['A11'].font = Font(bold=True)
    ws['B11'].font = Font(bold=True)

    pie = PieChart()
    labels = Reference(ws, min_col=1, min_row=12, max_row=15)
    data = Reference(ws, min_col=2, min_row=11, max_row=15)
    pie.add_data(data, titles_from_data=True)
    pie.set_categories(labels)
    pie.title = "Distribuição por Intenção"
    pie.dataLabels = DataLabelList()
    pie.dataLabels.showPercent = True
    ws.add_chart(pie, "D10")

    # --- Seção 3: Principais Recursos de SERP ---
    ws['A20'] = "Principais Recursos de SERP"
    ws['A20'].font = Font(size=14, bold=True)
    ws['A20'].fill = section_fill
    top_serp = sorted(serp_counts.items(), key=lambda x: x[1], reverse=True)[:5]
    ws.append(["Recurso", "Quantidade"])
    for feature, count in top_serp:
        ws.append([feature, count])
    for row in range(21, 27):
        for col in ['A', 'B']:
            ws[f'{col}{row}'].border = thin_border
    ws['A21'].font = Font(bold=True)
    ws['B21'].font = Font(bold=True)

    bar = BarChart()
    data = Reference(ws, min_col=2, min_row=21, max_row=26)
    cats = Reference(ws, min_col=1, min_row=22, max_row=26)
    bar.add_data(data, titles_from_data=True)
    bar.set_categories(cats)
    bar.title = "Top 5 Recursos de SERP"
    ws.add_chart(bar, "D20")

    # --- Seção 4: Jornada do Cliente ---
    ws['A30'] = "Distribuição por Jornada"
    ws['A30'].font = Font(size=14, bold=True)
    ws['A30'].fill = section_fill
    etapas = ["Conscientização", "Consideração", "Decisão", "Fidelização"]
    ws.append(["Etapa", "Quantidade"])
    for etapa in etapas:
        ws.append([etapa, jornada_counts.get(etapa, 0)])
    for row in range(31, 36):
        for col in ['A', 'B']:
            ws[f'{col}{row}'].border = thin_border
    ws['A31'].font = Font(bold=True)
    ws['B31'].font = Font(bold=True)

    pie_jornada = PieChart()
    labels_j = Reference(ws, min_col=1, min_row=32, max_row=35)
    data_j = Reference(ws, min_col=2, min_row=31, max_row=35)
    pie_jornada.add_data(data_j, titles_from_data=True)
    pie_jornada.set_categories(labels_j)
    pie_jornada.title = "Distribuição por Jornada"
    pie_jornada.dataLabels = DataLabelList()
    pie_jornada.dataLabels.showPercent = True
    ws.add_chart(pie_jornada, "D30")

    # --- Seção 5: CTR por Posição ---
    ws['A40'] = "CTR por Posição (Exemplo)"
    ws['A40'].font = Font(size=14, bold=True)
    ws['A40'].fill = section_fill
    if not ctr_export_df.empty:
        exemplo = ctr_export_df.iloc[0]
        ws['A41'] = f"Palavra: {exemplo['Keyword']}"
        ws.append(["Posição", "Cliques Mínimos", "Cliques Máximos"])
        for pos, (min_rate, max_rate) in ctr_rates.items():
            min_max = exemplo[f'Posicao {pos} ({int(min_rate*100)}%-{int(max_rate*100)}%)'].split('-')
            ws.append([f"Posição {pos}", int(min_max[0]), int(min_max[1])])
        for row in range(41, 53):
            for col in ['A', 'B', 'C']:
                ws[f'{col}{row}'].border = thin_border
        ws['A42'].font = Font(bold=True)
        ws['B42'].font = Font(bold=True)
        ws['C42'].font = Font(bold=True)

        line = LineChart()
        data_c = Reference(ws, min_col=2, min_row=42, max_col=3, max_row=52)
        cats_c = Reference(ws, min_col=1, min_row=43, max_row=52)
        line.add_data(data_c, titles_from_data=True)
        line.set_categories(cats_c)
        line.title = f"CTR para '{exemplo['Keyword']}'"
        line.y_axis.title = "Cliques"
        line.x_axis.title = "Posição"
        ws.add_chart(line, "D40")

    # --- Seção 6: Estratégia por Objetivo ---
    ws['A55'] = "Estratégia por Objetivo (Top 5)"
    ws['A55'].font = Font(size=14, bold=True)
    ws['A55'].fill = section_fill
    ws.append(["Palavra-chave", "Volume", "Estratégia", "Tipologia"])
    for _, row in estrategia_df.head(5).iterrows():
        ws.append([row["Palavra-chave"], row["Volume"], row["Estratégia"], row["Tipologia de Conteúdo"]])
    for row in range(56, 62):
        for col in ['A', 'B', 'C', 'D']:
            ws[f'{col}{row}'].border = thin_border
    for col in ['A', 'B', 'C', 'D']:
        ws[f'{col}56'].font = Font(bold=True)

    # --- Seção 7: Planejamento de Crescimento ---
    ws['A65'] = "Projeção de Crescimento"
    ws['A65'].font = Font(size=14, bold=True)
    ws['A65'].fill = section_fill
    volume_atual = calculo_df[calculo_df["Métrica"] == "Volume Atual (mensal)"]["Valor"].iloc[0]
    crescimento_mensal = calculo_df[calculo_df["Métrica"] == "Crescimento Desejado (%)"]["Valor"].iloc[0]
    meses_planejamento = int(calculo_df[calculo_df["Métrica"] == "Meses de Planejamento"]["Valor"].iloc[0])
    acessos = [volume_atual * (1 + crescimento_mensal / 100) ** i for i in range(meses_planejamento)]
    ws.append(["Mês", "Acessos Projetados"])
    for i, acesso in enumerate(acessos, start=1):
        ws.append([f"Mês {i}", round(acesso)])
    for row in range(66, 66 + meses_planejamento + 1):
        for col in ['A', 'B']:
            ws[f'{col}{row}'].border = thin_border
    ws['A66'].font = Font(bold=True)
    ws['B66'].font = Font(bold=True)

    line_c = LineChart()
    data_c = Reference(ws, min_col=2, min_row=66, max_row=66 + meses_planejamento)
    cats_c = Reference(ws, min_col=1, min_row=67, max_row=66 + meses_planejamento)
    line_c.add_data(data_c, titles_from_data=True)
    line_c.set_categories(cats_c)
    line_c.title = "Projeção de Crescimento"
    line_c.y_axis.title = "Acessos Mensais"
    line_c.x_axis.title = "Meses"
    ws.add_chart(line_c, "D65")

    # Ajustar largura das colunas
    adjust_column_width(ws)

    # Salvar o dashboard
    wb.save(os.path.join(folder_name, "Dashboard.xlsx"))
    print_status("Dashboard.xlsx gerado com sucesso!")

# =============================================================================
# Configuração Inicial e Criação da Pasta de Saída
# =============================================================================

print_status("Bem-vindo ao Script de Análise de Palavras-Chave para SEO!")
project_name = input("[PERGUNTA] Qual o nome do projeto? ").strip()
now = datetime.now()
folder_name = f"{project_name} {now.strftime('%d-%m-%Y')} {now.strftime('%H')} horas {now.strftime('%M')} minutos {now.strftime('%S')} segundos"
os.makedirs(folder_name, exist_ok=True)
print_status(f"Pasta de saída criada: {folder_name}")

if not os.path.exists("cidades_brasil.xlsx"):
    print_status("Erro: Arquivo 'cidades_brasil.xlsx' não encontrado!")
    raise FileNotFoundError("Arquivo 'cidades_brasil.xlsx' necessário")
if not os.path.exists("kw_negativas.docx"):
    print_status("Erro: Arquivo 'kw_negativas.docx' não encontrado!")
    raise FileNotFoundError("Arquivo 'kw_negativas.docx' necessário")

use_gpt = input("[PERGUNTA] Deseja conectar à API do ChatGPT para assistência? (s/n): ").lower()
if use_gpt == 's':
    print_status("A opção de API foi escolhida, mas este código usará o mapeamento interno para tipologia.")
objective = input(
    "[PERGUNTA] Qual o objetivo estratégico da análise?\n"
    "Opções: 1) Captura de leads, 2) Vendas no e-commerce, 3) Mais acessos, 4) Monetização com Adsense, 5) Branding/Autoridade, 6) Outro\n"
    "Digite o número ou descreva: "
)
volume_atual = int(input("[PERGUNTA] Qual o volume de acessos mensal atual do site? "))
crescimento_mensal = float(input("[PERGUNTA] Qual o percentual de crescimento desejado por mês? (ex: 10 para 10%): "))
meses_planejamento = int(input("[PERGUNTA] Quantos meses será o planejamento? "))
palavras_por_mes = int(input("[PERGUNTA] Quantas palavras-chave serão trabalhadas por mês? "))

folder_path = os.getcwd()
print_status(f"Usando a pasta atual como fonte das planilhas: {folder_path}")

print_status("Carregando lista de cidades do Brasil do arquivo 'cidades_brasil.xlsx'...")
cidades_brasil = carregar_cidades_brasil()
print_status(f"{len(cidades_brasil)} cidades carregadas para exclusão.")

# =============================================================================
# Fase 1 – Aglutinar as Planilhas
# =============================================================================

print_status("Iniciando Fase 1: Aglutinando planilhas...")
all_data = []
for filename in os.listdir(folder_path):
    if filename.endswith('.xlsx') and not filename.startswith(project_name) and filename != "cidades_brasil.xlsx":
        print_status(f"Lendo arquivo: {filename}")
        try:
            df = pd.read_excel(os.path.join(folder_path, filename))
            all_data.append(df)
        except Exception as e:
            print_status(f"Erro ao ler {filename}: {str(e)}")
            raise
if not all_data:
    print_status("Erro: Nenhuma planilha .xlsx encontrada na pasta (exceto 'cidades_brasil.xlsx')!")
    raise ValueError("Nenhum arquivo válido encontrado")

try:
    combined_df = pd.concat(all_data, ignore_index=True)
    combined_df = combined_df.sort_values(by=['Keyword'], ascending=True)
except Exception as e:
    print_status(f"Erro ao concatenar planilhas: {str(e)}")
    raise

volume_col = None
for col in combined_df.columns:
    if "volume" in col.lower():
        volume_col = col
        break
if volume_col:
    combined_df = combined_df.sort_values(by=[volume_col], ascending=False)

wb = Workbook()
ws = wb.active
ws.title = "Visao Geral de Palavras"
for r in dataframe_to_rows(combined_df, index=False, header=True):
    ws.append(r)
volume_col_index = get_column_index(ws, volume_col) if volume_col else None
if volume_col_index and not combined_df[volume_col].dropna().empty:
    apply_heatmap(ws, volume_col_index, combined_df[volume_col])
apply_header_style(ws)
apply_content_style(ws)
adjust_column_width(ws)
visao_geral_filename = os.path.join(folder_name, "Visao Geral de Palavras.xlsx")
wb.save(visao_geral_filename)
print_status("Fase 1 concluída: Planilha 'Visao Geral de Palavras.xlsx' gerada!")

print_status("Formatando a planilha temporária 'combined_df_temp.xlsx'...")
wb_temp = Workbook()
ws_temp = wb_temp.active
ws_temp.title = "Dados Combinados Temporários"
for r in dataframe_to_rows(combined_df, index=False, header=True):
    ws_temp.append(r)
volume_col_index = get_column_index(ws_temp, volume_col) if volume_col else None
if volume_col_index and not combined_df[volume_col].dropna().empty:
    apply_heatmap(ws_temp, volume_col_index, combined_df[volume_col])
apply_header_style(ws_temp)
apply_content_style(ws_temp)
adjust_column_width(ws_temp)
wb_temp.save(os.path.join(folder_name, "combined_df_temp.xlsx"))
print_status("Planilha 'combined_df_temp.xlsx' formatada com sucesso!")

plt.figure(figsize=(8, 4))
top_10 = combined_df.head(10)
bars = plt.bar(top_10['Keyword'], top_10[volume_col])
plt.title("Top 10 Palavras por Volume de Busca")
plt.xlabel("Palavras-chave")
plt.ylabel("Volume de Busca")
plt.xticks(rotation=45)
for bar in bars:
    yval = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2, yval, int(yval), ha='center', va='bottom')
plt.tight_layout()
plt.savefig(os.path.join(folder_name, "visao_geral.png"))
plt.close()

# =============================================================================
# Fase 2 – Separação por Intent
# =============================================================================

print_status("Iniciando Fase 2: Separando por intenção de busca...")
intents = ['Informational', 'Transactional', 'Commercial', 'Navigational']
wb_intent = Workbook()
ws_overview = wb_intent.active
ws_overview.title = "Visao Geral"
for r in dataframe_to_rows(combined_df, index=False, header=True):
    ws_overview.append(r)
volume_col_index = get_column_index(ws_overview, volume_col) if volume_col else None
if volume_col_index and not combined_df[volume_col].dropna().empty:
    apply_heatmap(ws_overview, volume_col_index, combined_df[volume_col])
apply_header_style(ws_overview)
apply_content_style(ws_overview)
adjust_column_width(ws_overview)

intent_counts = {}
for intent in intents:
    print_status(f"Criando aba para Intent: {intent}")
    try:
        intent_df = combined_df[combined_df['Intent'].str.contains(intent, case=False, na=False)].sort_values(by=['Keyword'], ascending=True)
        if volume_col:
            intent_df = intent_df.sort_values(by=[volume_col], ascending=False)
        intent_counts[intent] = len(intent_df)
        ws_intent = wb_intent.create_sheet(intent)
        for r in dataframe_to_rows(intent_df, index=False, header=True):
            ws_intent.append(r)
        volume_col_index = get_column_index(ws_intent, volume_col) if volume_col else None
        if volume_col_index and not intent_df[volume_col].dropna().empty:
            apply_heatmap(ws_intent, volume_col_index, intent_df[volume_col])
        apply_header_style(ws_intent)
        apply_content_style(ws_intent)
        adjust_column_width(ws_intent)
    except Exception as e:
        print_status(f"Erro ao processar Intent '{intent}': {str(e)}")
        raise

no_intent_df = combined_df[combined_df['Intent'].isna()].sort_values(by=['Keyword'], ascending=True)
if volume_col:
    no_intent_df = no_intent_df.sort_values(by=[volume_col], ascending=False)
ws_no_intent = wb_intent.create_sheet("Sem Intent")
for r in dataframe_to_rows(no_intent_df, index=False, header=True):
    ws_no_intent.append(r)
volume_col_index = get_column_index(ws_no_intent, volume_col) if volume_col else None
if volume_col_index and not no_intent_df[volume_col].dropna().empty:
    apply_heatmap(ws_no_intent, volume_col_index, no_intent_df[volume_col])
apply_header_style(ws_no_intent)
apply_content_style(ws_no_intent)
adjust_column_width(ws_no_intent)

intents_filename = os.path.join(folder_name, "Intents.xlsx")
wb_intent.save(intents_filename)
print_status("Fase 2 concluída: Planilha 'Intents.xlsx' gerada!")

plt.figure(figsize=(6, 6))
intent_values = [intent_counts.get(i, 0) for i in intents]
plt.pie(intent_values, labels=intents, autopct='%1.1f%%', colors=['#FF9999', '#66B2FF', '#99FF99', '#FFCC99'])
plt.title("Distribuição por Intenção de Busca")
for i, (value, label) in enumerate(zip(intent_values, intents)):
    angle = sum(intent_values[:i]) + value / 2
    angle_rad = angle * 2 * np.pi / sum(intent_values)
    x = 0.5 * np.cos(angle_rad)
    y = 0.5 * np.sin(angle_rad)
    plt.text(x, y, str(value), ha='center', va='center')
plt.savefig(os.path.join(folder_name, "intents.png"))
plt.close()

# =============================================================================
# Fase 3 – Separação por SERP Features
# =============================================================================

print_status("Iniciando Fase 3: Separando por SERP Features...")
wb_serp = Workbook()
ws_serp_overview = wb_serp.active
ws_serp_overview.title = "Visao Geral"
for r in dataframe_to_rows(combined_df, index=False, header=True):
    ws_serp_overview.append(r)
volume_col_index = get_column_index(ws_serp_overview, volume_col) if volume_col else None
if volume_col_index and not combined_df[volume_col].dropna().empty:
    apply_heatmap(ws_serp_overview, volume_col_index, combined_df[volume_col])
apply_header_style(ws_serp_overview)
apply_content_style(ws_serp_overview)
adjust_column_width(ws_serp_overview)

serp_col = None
for col in combined_df.columns:
    if col.strip().lower() == "serp features":
        serp_col = col
        break

serp_counts = {}
if serp_col:
    features_set = set()
    for val in combined_df[serp_col].dropna():
        for token in str(val).split(','):
            token = token.strip()
            if token:
                features_set.add(token)
    for feature in features_set:
        print_status(f"Criando aba para SERP Feature: {feature}")
        feature_df = combined_df[combined_df[serp_col].str.contains(feature, case=False, na=False)]
        if volume_col:
            feature_df = feature_df.sort_values(by=[volume_col], ascending=False)
        serp_counts[feature] = len(feature_df)
        if not feature_df.empty:
            ws_feature = wb_serp.create_sheet(feature)
            for r in dataframe_to_rows(feature_df, index=False, header=True):
                ws_feature.append(r)
            volume_col_index = get_column_index(ws_feature, volume_col) if volume_col else None
            if volume_col_index and not feature_df[volume_col].dropna().empty:
                apply_heatmap(ws_feature, volume_col_index, feature_df[volume_col])
            apply_header_style(ws_feature)
            apply_content_style(ws_feature)
            adjust_column_width(ws_feature)
else:
    print_status("Aviso: Coluna 'SERP Features' não encontrada. Pulando separação por SERP Features.")
serp_features_filename = os.path.join(folder_name, "SERP Features.xlsx")
wb_serp.save(serp_features_filename)
print_status("Fase 3 concluída: Planilha 'SERP Features.xlsx' gerada!")

top_serp = sorted(serp_counts.items(), key=lambda x: x[1], reverse=True)[:4]
plt.figure(figsize=(8, 4))
bars = plt.bar([x[0] for x in top_serp], [x[1] for x in top_serp])
plt.title("Principais Recursos de SERP")
plt.xlabel("Recurso")
plt.ylabel("Quantidade")
for bar in bars:
    yval = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2, yval, int(yval), ha='center', va='bottom')
plt.savefig(os.path.join(folder_name, "serp_features.png"))
plt.close()

# =============================================================================
# Fase 4 – Mapeamento por Jornada e Tipologia
# =============================================================================

print_status("Iniciando Fase 4: Mapeando por Jornada e Tipologia...")
jornada_list = []
tipologia_list = []
for idx, row in combined_df.iterrows():
    jornada_list.append(get_etapa_da_jornada(row.get('Intent', '')))
    tipologia_list.append(get_tipologia_sugerida(row))
combined_df['Etapa da Jornada'] = jornada_list
combined_df['Tipologia Sugerida'] = tipologia_list

cols_to_drop = ["CPC (USD)", "Competitive Density", "Number of Results"]
combined_df = combined_df.drop(columns=cols_to_drop, errors='ignore')

if volume_col:
    combined_df = combined_df.sort_values(by=[volume_col], ascending=False)
else:
    combined_df = combined_df.sort_values(by=['Keyword'], ascending=True)

wb_journey = Workbook()
ws_journey_overview = wb_journey.active
ws_journey_overview.title = "Visao Geral da Jornada"
for r in dataframe_to_rows(combined_df, index=False, header=True):
    ws_journey_overview.append(r)
volume_col_index = get_column_index(ws_journey_overview, volume_col) if volume_col else None
if volume_col_index and not combined_df[volume_col].dropna().empty:
    apply_heatmap(ws_journey_overview, volume_col_index, combined_df[volume_col])
apply_header_style(ws_journey_overview)
apply_content_style(ws_journey_overview)
adjust_column_width(ws_journey_overview)

etapas = ["Conscientização", "Consideração", "Decisão", "Fidelização", "Sem Jornada Definida"]
jornada_counts = {}
for etapa in etapas:
    print_status(f"Criando aba para Jornada: {etapa}")
    etapa_df = combined_df[combined_df['Etapa da Jornada'] == etapa]
    if volume_col:
        etapa_df = etapa_df.sort_values(by=[volume_col], ascending=False)
    jornada_counts[etapa] = len(etapa_df)
    ws_etapa = wb_journey.create_sheet(etapa)
    for r in dataframe_to_rows(etapa_df, index=False, header=True):
        ws_etapa.append(r)
    volume_col_index = get_column_index(ws_etapa, volume_col) if volume_col else None
    if volume_col_index and not etapa_df[volume_col].dropna().empty:
        apply_heatmap(ws_etapa, volume_col_index, etapa_df[volume_col])
    apply_header_style(ws_etapa)
    apply_content_style(ws_etapa)
    adjust_column_width(ws_etapa)

jornada_filename = os.path.join(folder_name, "Jornada e Tipologias.xlsx")
wb_journey.save(jornada_filename)
print_status("Fase 4 concluída: Planilha 'Jornada e Tipologias.xlsx' gerada!")

plt.figure(figsize=(6, 6))
jornada_values = [jornada_counts[e] for e in etapas if e in jornada_counts]
plt.pie(jornada_values, labels=[e for e in etapas if e in jornada_counts], autopct='%1.1f%%', colors=['#FF6666', '#FFCC66', '#66CCFF', '#66FF66', '#999999'])
plt.title("Distribuição por Etapa da Jornada")
for i, (value, label) in enumerate(zip(jornada_values, [e for e in etapas if e in jornada_counts])):
    angle = sum(jornada_values[:i]) + value / 2
    angle_rad = angle * 2 * np.pi / sum(jornada_values)
    x = 0.5 * np.cos(angle_rad)
    y = 0.5 * np.sin(angle_rad)
    plt.text(x, y, str(value), ha='center', va='center')
plt.savefig(os.path.join(folder_name, "jornada.png"))
plt.close()

## Fase 5 – CTR por Posição
print_status("Iniciando Fase 5: Calculando CTR por posição...")
ctr_rates = {
    1: (0.25, 0.35), 2: (0.15, 0.20), 3: (0.10, 0.15), 4: (0.07, 0.10), 5: (0.05, 0.07),
    6: (0.04, 0.06), 7: (0.03, 0.05), 8: (0.02, 0.04), 9: (0.02, 0.03), 10: (0.01, 0.02)
}
if volume_col:
    ctr_df = combined_df[(combined_df[volume_col] > 0) & (combined_df[volume_col].notna())].copy()
    for pos, (min_rate, max_rate) in ctr_rates.items():
        ctr_df[f'Posicao {pos}'] = ctr_df[volume_col].apply(lambda x: f"{int(x * min_rate)} - {int(x * max_rate)}")
    selected_columns = ['Keyword', volume_col, 'Intent', 'Trend'] + [f'Posicao {i}' for i in range(1, 11)]
    ctr_export_df = ctr_df[selected_columns].copy()
    for pos, (min_rate, max_rate) in ctr_rates.items():
        old_col = f'Posicao {pos}'
        new_col = f'Posicao {pos} ({int(min_rate*100)}%-{int(max_rate*100)}%)'
        ctr_export_df.rename(columns={old_col: new_col}, inplace=True)
    wb_ctr = Workbook()
    ws_ctr = wb_ctr.active
    ws_ctr.title = "CTR por Posicao"
    for r in dataframe_to_rows(ctr_export_df, index=False, header=True):
        ws_ctr.append(r)
    volume_col_index = get_column_index(ws_ctr, volume_col)
    if volume_col_index and not ctr_export_df[volume_col].dropna().empty:
        apply_heatmap(ws_ctr, volume_col_index, ctr_export_df[volume_col])
    apply_header_style(ws_ctr)
    apply_content_style(ws_ctr)
    adjust_column_width(ws_ctr)
    ctr_filename = os.path.join(folder_name, "CTR por Posicao.xlsx")
    wb_ctr.save(ctr_filename)
    print_status("Fase 5 concluída: Planilha 'CTR por Posicao.xlsx' gerada!")
else:
    print_status("Aviso: Nenhuma coluna de volume encontrada. Pulando Fase 5.")
    ctr_export_df = pd.DataFrame()

if not ctr_export_df.empty:
    exemplo_ctr = ctr_export_df.iloc[0]
    plt.figure(figsize=(8, 4))
    positions = range(1, 11)
    ctr_min = [float(exemplo_ctr[f'Posicao {i} ({int(min_rate*100)}%-{int(max_rate*100)}%)'].split('-')[0]) for i, (min_rate, max_rate) in ctr_rates.items()]
    ctr_max = [float(exemplo_ctr[f'Posicao {i} ({int(min_rate*100)}%-{int(max_rate*100)}%)'].split('-')[1]) for i, (min_rate, max_rate) in ctr_rates.items()]
    plt.plot(positions, ctr_min, label="Cenário Pessimista", marker='o')
    plt.plot(positions, ctr_max, label="Cenário Otimista", marker='o')
    plt.title(f"Estimativa de CTR por Posição ({exemplo_ctr['Keyword']})")
    plt.xlabel("Posição")
    plt.ylabel("Cliques")
    for i, (min_val, max_val) in enumerate(zip(ctr_min, ctr_max)):
        plt.text(positions[i], min_val, int(min_val), ha='center', va='bottom')
        plt.text(positions[i], max_val, int(max_val), ha='center', va='bottom')
    plt.legend()
    plt.savefig(os.path.join(folder_name, "ctr_posicao.png"))
    plt.close()

# =============================================================================
# Fase 6 – Estratégia por Objetivo com Palavras-Chave
# =============================================================================

print_status("Iniciando Fase 6: Gerando estratégias por objetivo com palavras-chave...")
estrategia_df, objetivo_selecionado = criar_planilha_palavras_por_estrategia(folder_name, objective, combined_df)
print_status("Fase 6 concluída: Planilha 'Palavras por Estratégia.xlsx' gerada!")

# =============================================================================
# Fase 7 – Planejamento de Crescimento
# =============================================================================

print_status("Iniciando Fase 7: Gerando planejamento de crescimento...")
result = criar_planilha_planejamento_crescimento(folder_name, combined_df, volume_atual, crescimento_mensal, meses_planejamento, palavras_por_mes, objective, cidades_brasil)
if result is None:
    print_status("Erro na Fase 7. Abortando execução.")
    raise ValueError("Fase 7 falhou devido à ausência de coluna de volume ou outro erro.")
calculo_df, palavras_selecionadas, cauda_curta, cauda_media, cauda_longa, palavras_semantico, palavras_blog, meses, colunas_selecao = result
print_status("Fase 7 concluída: Planilha 'Planejamento de Crescimento.xlsx' gerada!")

plt.figure(figsize=(8, 4))
meses_grafico = [f'Mês {i+1}' for i in range(meses_planejamento)]
acessos = [volume_atual * (1 + crescimento_mensal / 100) ** i for i in range(meses_planejamento)]
plt.plot(meses_grafico, acessos, marker='o')
plt.title("Projeção de Crescimento")
plt.xlabel("Meses")
plt.ylabel("Acessos Mensais")
for i, val in enumerate(acessos):
    plt.text(i, val, int(val), ha='center', va='bottom')
plt.savefig(os.path.join(folder_name, "crescimento.png"))
plt.close()

# =============================================================================
# Fase 8 – Top 100 Palavras por Tipo
# =============================================================================

print_status("Iniciando Fase 8: Gerando top 100 palavras por tipo...")
top_palavras_por_tipo = criar_planilha_top_palavras_por_tipo(folder_name, combined_df)
print_status("Fase 8 concluída: Planilha 'Top 100 Palavras por Tipo.xlsx' gerada!")

# =============================================================================
# Fase 8.5 – Palavras para Ads Filtradas
# =============================================================================
print_status("Iniciando Fase 8.5: Gerando palavras para Ads Filtradas...")
palavras_ads_filtradas, palavras_excluidas = criar_planilha_palavras_para_ads_filtradas(folder_name, combined_df)
print_status("Fase 8.5 concluída: Planilha 'Palavras para Ads Filtradas.xlsx' gerada!")

# =============================================================================
# Fase 9 – Geração do Dashboard Profissional
print_status("Iniciando Fase 9: Gerando Dashboard Profissional...")
criar_dashboard_profissional(folder_name, combined_df, intent_counts, serp_counts, jornada_counts, ctr_export_df, estrategia_df, calculo_df, meses, volume_col, objective, now, ctr_rates)
print_status("Fase 9 concluída: Dashboard.xlsx gerado!")

def dict_to_xml(tag, d):
    elem = ET.Element(tag)
    for key, val in d.items():
        # Garantir que a chave seja uma string válida para XML
        key = re.sub(r'[^a-zA-Z0-9_]', '_', str(key))  # Substitui caracteres inválidos por '_'
        
        if isinstance(val, dict):
            child = dict_to_xml(key, val)
            elem.append(child)
        elif isinstance(val, list):
            for i, item in enumerate(val):
                if isinstance(item, dict):
                    child = dict_to_xml(key, item)
                else:
                    child = ET.Element(key)
                    # Sanitizar completamente o texto
                    safe_text = str(item)
                    # Remover caracteres de controle e quaisquer caracteres inválidos para XML
                    safe_text = ''.join(c for c in safe_text if 32 <= ord(c) <= 0x10FFFF and c not in '\x00-\x08\x0B\x0C\x0E-\x1F\x7F')
                    # Escapar caracteres especiais XML
                    safe_text = safe_text.encode('utf-8', errors='xmlcharrefreplace').decode('utf-8')
                    child.text = safe_text if safe_text.strip() else ' '  # Garante texto não-vazio
                elem.append(child)
        else:
            child = ET.Element(key)
            # Sanitizar completamente o texto
            safe_text = str(val)
            # Remover caracteres de controle e quaisquer caracteres inválidos para XML
            safe_text = ''.join(c for c in safe_text if 32 <= ord(c) <= 0x10FFFF and c not in '\x00-\x08\x0B\x0C\x0E-\x1F\x7F')
            # Escapar caracteres especiais XML
            safe_text = safe_text.encode('utf-8', errors='xmlcharrefreplace').decode('utf-8')
            child.text = safe_text if safe_text.strip() else ' '  # Garante texto não-vazio
            elem.append(child)
    return elem    
# Configuração Inicial e Criação da Pasta de Saída
# =============================================================================
# Geração do Relatório Analítico Detalhado
# =============================================================================

print_status("Gerando Relatório Analítico Detalhado.docx...")
doc = Document()

add_title(doc, f"Relatório Analítico Detalhado - Projeto {project_name}")

add_subtitle(doc, "Introdução")
add_paragraph(doc, f"Este relatório apresenta uma análise detalhada e fundamentada das palavras-chave fornecidas para o projeto {project_name}, com o objetivo de otimizar a estratégia de SEO do site. Foram realizadas nove fases analíticas, cada uma com propósitos específicos para compreender o comportamento de busca, o potencial de tráfego e as oportunidades de conteúdo. Abaixo, detalhamos cada fase, os métodos utilizados, os resultados obtidos e recomendações estratégicas.")

add_subtitle(doc, "Fase 1: Visão Geral de Palavras")
add_paragraph(doc, "Objetivo: Consolidar todas as palavras-chave de diferentes fontes em uma única planilha para fornecer uma visão geral do volume de busca, intenção e características de SERP.")
add_paragraph(doc, "Método: As planilhas foram lidas, concatenadas usando pandas e ordenadas pelo volume de busca (quando disponível). Uma versão temporária foi salva para depuração.")
add_paragraph(doc, f"Resultado: Foram analisadas {len(combined_df)} palavras-chave únicas. O gráfico abaixo destaca as 10 principais por volume de busca.")
add_image(doc, os.path.join(folder_name, "visao_geral.png"))
add_paragraph(doc, "Recomendações: Priorizar palavras de alto volume para estratégias de curto prazo e explorar termos de cauda longa para ganhos sustentáveis.")

add_subtitle(doc, "Fase 2: Separação por Intenção de Busca")
add_paragraph(doc, "Objetivo: Classificar as palavras-chave em intenções de busca (Informacional, Transacional, Comercial, Navegacional) para alinhar o conteúdo às expectativas dos usuários.")
add_paragraph(doc, "Método: Filtragem baseada na coluna 'Intent' com ordenação por volume.")
add_paragraph(doc, f"Resultado: Distribuição das intenções: {', '.join([f'{intent}: {intent_counts.get(intent, 0)}' for intent in intents])}. Veja o gráfico abaixo.")
add_image(doc, os.path.join(folder_name, "intents.png"))
add_paragraph(doc, "Recomendações: Criar conteúdo específico para cada intenção, como guias para Informacional e páginas de produto para Transacional.")

add_subtitle(doc, "Fase 3: Separação por Recursos de SERP")
add_paragraph(doc, "Objetivo: Identificar palavras-chave associadas a recursos de SERP para explorar oportunidades de visibilidade.")
add_paragraph(doc, "Método: Extração e contagem de features da coluna 'SERP Features', com separação em abas.")
add_paragraph(doc, f"Resultado: Principais recursos encontrados: {', '.join([f'{feat}: {count}' for feat, count in top_serp])}. Veja o gráfico.")
add_image(doc, os.path.join(folder_name, "serp_features.png"))
add_paragraph(doc, "Recomendações: Otimizar para Featured Snippets e Local Pack quando aplicável, aumentando CTR.")

add_subtitle(doc, "Fase 4: Mapeamento por Jornada e Tipologia")
add_paragraph(doc, "Objetivo: Mapear palavras-chave às etapas da jornada do cliente e sugerir tipologias de conteúdo.")
add_paragraph(doc, "Método: Uso de funções personalizadas para classificar intenções em etapas e sugerir tipologias baseadas em SERP Features.")
add_paragraph(doc, f"Resultado: Distribuição por etapa: {', '.join([f'{etapa}: {jornada_counts.get(etapa, 0)}' for etapa in etapas])}. Veja o gráfico.")
add_image(doc, os.path.join(folder_name, "jornada.png"))
add_paragraph(doc, "Recomendações: Desenvolver funis de conteúdo alinhados à jornada, como blogs para Conscientização e comparativos para Consideração.")

add_subtitle(doc, "Fase 5: CTR por Posição")
add_paragraph(doc, "Objetivo: Estimar cliques potenciais por posição no ranking para cada palavra-chave.")
add_paragraph(doc, "Método: Aplicação de taxas de CTR padrão por posição ao volume de busca.")
if not ctr_export_df.empty:
    add_paragraph(doc, f"Resultado: Exemplo para '{exemplo_ctr['Keyword']}' (volume {exemplo_ctr[volume_col]}). Veja a estimativa abaixo.")
    add_image(doc, os.path.join(folder_name, "ctr_posicao.png"))
    add_paragraph(doc, "Recomendações: Focar em alcançar as primeiras posições para palavras de alto volume.")
else:
    add_paragraph(doc, "Resultado: Análise não realizada devido à ausência de dados de volume.")

add_subtitle(doc, "Fase 6: Estratégia por Objetivo com Palavras-Chave")
add_paragraph(doc, "Objetivo: Associar palavras-chave a estratégias específicas conforme o objetivo selecionado.")
add_paragraph(doc, "Método: Mapeamento de intenções a objetivos e fusão com estratégias predefinidas.")
add_paragraph(doc, f"Resultado: Para o objetivo '{objetivo_selecionado}', exemplo: '{estrategia_df.iloc[0]['Palavra-chave']}' (volume {estrategia_df.iloc[0]['Volume']}), Estratégia: '{estrategia_df.iloc[0]['Estratégia']}'.")
add_paragraph(doc, "Recomendações: Implementar as tipologias sugeridas para maximizar o impacto do objetivo escolhido.")

add_subtitle(doc, "Fase 7: Planejamento de Crescimento")
add_paragraph(doc, "Objetivo: Projetar o crescimento de tráfego com base em volume atual, meta de crescimento e palavras-chave selecionadas.")
add_paragraph(doc, "Método: Cálculo de metas com exclusão de termos geográficos e segmentação por cauda.")
add_paragraph(doc, f"Resultado: Projeção para {meses_planejamento} meses, volume inicial {volume_atual}, crescimento {crescimento_mensal}% ao mês. Veja o gráfico.")
add_image(doc, os.path.join(folder_name, "crescimento.png"))
add_paragraph(doc, "Recomendações: Priorizar palavras selecionadas na aba 'Seleção de Palavras' e monitorar o progresso mensal.")

add_subtitle(doc, "Fase 8: Top 100 Palavras por Tipo")
add_paragraph(doc, "Objetivo: Identificar as 100 melhores palavras-chave por tipo com base no volume de busca.")
add_paragraph(doc, "Método: Agrupamento por uma coluna de tipo (se disponível) ou clustering automático das palavras-chave.")
if top_palavras_por_tipo:
    primeiro_tipo = list(top_palavras_por_tipo.keys())[0]
    exemplo_df = top_palavras_por_tipo[primeiro_tipo]
    add_paragraph(doc, f"Resultado: Tipos analisados: {', '.join(top_palavras_por_tipo.keys())}. Exemplo para '{primeiro_tipo}': '{exemplo_df.iloc[0]['Keyword']}' (volume {exemplo_df.iloc[0][volume_col]}).")
    add_paragraph(doc, "Recomendações: Focar em palavras de alto volume por tipo para campanhas segmentadas.")
else:
    add_paragraph(doc, "Resultado: Nenhum tipo identificado devido a dados insuficientes.")

# Adicionar a Fase 8.5 aqui
add_subtitle(doc, "Fase 8.5: Palavras para Ads Filtradas")
add_paragraph(doc, "Objetivo: Filtrar palavras-chave adequadas para campanhas de anúncios, excluindo termos negativos listados em 'kw_negativas.docx'.")
add_paragraph(doc, "Método: Leitura de palavras negativas de um arquivo Word, exclusão de palavras contendo esses termos e separação em duas abas: 'Palavras Filtradas' e 'Palavras Excluídas'.")
add_paragraph(doc, f"Resultado: {len(palavras_ads_filtradas)} palavras filtradas e {len(palavras_excluidas)} excluídas.")
add_paragraph(doc, "Recomendações: Usar as palavras filtradas para campanhas de anúncios e revisar as excluídas para ajustes na lista de negativas.")

add_subtitle(doc, "Fase 9: Dashboard Profissional")
add_paragraph(doc, "Objetivo: Criar um dashboard interativo no Excel para visualização consolidada dos principais insights.")
add_paragraph(doc, "Método: Geração de tabelas e gráficos (pizza, barras e linhas) com dados de intenções, SERP, jornada, CTR e crescimento.")
add_paragraph(doc, "Resultado: Dashboard gerado em 'Dashboard.xlsx', contendo resumo geral, distribuições e projeções.")
add_paragraph(doc, "Recomendações: Utilizar o dashboard para apresentações e monitoramento estratégico.")

add_subtitle(doc, "Conclusão e Recomendações Finais")
add_paragraph(doc, f"A análise do projeto {project_name} oferece insights estratégicos para otimizar o SEO. Recomendamos: (1) Priorizar palavras de alto volume e baixa concorrência, (2) Implementar tipologias de conteúdo sugeridas, (3) Seguir o planejamento de crescimento para atingir as metas de tráfego, e (4) Monitorar os resultados regularmente com o dashboard gerado.")

doc.save(os.path.join(folder_name, "Relatório Analítico Detalhado.docx"))
print_status("Relatório Analítico Detalhado.docx gerado com sucesso na pasta " + folder_name)

# Geração do XML
print_status("Gerando resultados_finais.xml...")
resultados = {
    "Projeto": project_name,
    "Data": now.strftime('%d-%m-%Y %H:%M:%S'),
    "Objetivo": objective,
    "Fase1": {
        "TotalPalavras": len(combined_df),
        "VolumeMedio": round(combined_df[volume_col].mean(), 2) if volume_col and not combined_df[volume_col].dropna().empty else "N/A"
    },
    "Fase2": intent_counts,
    "Fase3": serp_counts,
    "Fase4": jornada_counts,
    "Fase5": {
        "ExemploCTR": {
            "Keyword": exemplo_ctr['Keyword'],
            "Volume": exemplo_ctr[volume_col] if volume_col else "N/A"
        } if not ctr_export_df.empty else "Nenhum dado"
    },
    "Fase6": {
        "ObjetivoSelecionado": objetivo_selecionado,
        "Exemplo": {
            "Palavra": estrategia_df.iloc[0]['Palavra-chave'],
            "Volume": estrategia_df.iloc[0]['Volume'],
            "Estrategia": estrategia_df.iloc[0]['Estratégia']
        }
    },
    "Fase7": {
        "VolumeAtual": volume_atual,
        "CrescimentoMensal": crescimento_mensal,
        "MesesPlanejamento": meses_planejamento
    },
    "Fase8": {
        "Tipos": {tipo: len(df) for tipo, df in top_palavras_por_tipo.items()}
    },
    "Fase8_5": {
        "PalavrasFiltradas": len(palavras_ads_filtradas),
        "PalavrasExcluidas": len(palavras_excluidas)
    },
    "Fase9": "Dashboard gerado"
}

root = dict_to_xml("Resultados", resultados)
xml_str = ET.tostring(root, encoding='utf-8', method='xml')

try:
    pretty_xml = minidom.parseString(xml_str).toprettyxml(indent="  ")
except Exception as e:
    print_status(f"Erro ao formatar XML: {str(e)}. Salvando versão bruta.")
    pretty_xml = xml_str.decode('utf-8')

with open(os.path.join(folder_name, "resultados_finais.xml"), "w", encoding="utf-8") as xml_file:
    xml_file.write(pretty_xml)
print_status("resultados_finais.xml gerado com sucesso na pasta " + folder_name)
print_status("Obrigado por usar o script do Consultor SEO Anderson Melo!")
